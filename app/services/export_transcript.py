"""Экспорт расшифровки по ролям в человекочитаемые форматы: txt и docx.

В отличие от субтитров (export_subs), тут не таймкоды для плеера, а аккуратный
конспект урока: кто говорил и что сказал, с таймкодом-меткой.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any


def _stamp(seconds: float) -> str:
    total = int(max(0.0, seconds))
    return f"{total // 60:02d}:{total % 60:02d}"


def to_txt(timeline: list[dict[str, Any]], title: str = "Расшифровка урока") -> str:
    """Простой текст: заголовок и строки «[MM:SS] Роль: текст»."""
    lines = [title, ""]
    for utt in timeline:
        name = utt.get("display_name", utt.get("source_speaker", "?"))
        text = str(utt.get("text", "")).strip()
        lines.append(f"[{_stamp(float(utt['start']))}] {name}: {text}")
    return "\n".join(lines) + "\n"


def _build_document(timeline: list[dict[str, Any]], title: str):
    """Собрать docx-документ: заголовок + реплики (метка и роль жирным)."""
    from docx import Document  # ленивый импорт: python-docx нужен только тут

    document = Document()
    document.add_heading(title, level=1)
    for utt in timeline:
        name = utt.get("display_name", utt.get("source_speaker", "?"))
        text = str(utt.get("text", "")).strip()
        paragraph = document.add_paragraph()
        head = paragraph.add_run(f"[{_stamp(float(utt['start']))}] {name}: ")
        head.bold = True
        paragraph.add_run(text)
    return document


def write_docx(
    timeline: list[dict[str, Any]],
    path: str | Path,
    title: str = "Расшифровка урока",
) -> Path:
    """Сохранить .docx по пути."""
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    _build_document(timeline, title).save(str(out))
    return out


def docx_bytes(
    timeline: list[dict[str, Any]],
    title: str = "Расшифровка урока",
) -> bytes:
    """Собрать .docx в память (для отдачи из API)."""
    import io

    buffer = io.BytesIO()
    _build_document(timeline, title).save(buffer)
    return buffer.getvalue()
